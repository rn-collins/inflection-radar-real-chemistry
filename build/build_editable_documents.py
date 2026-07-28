from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
ASSETS.mkdir(exist_ok=True)

INK = "10242B"; TEAL = "087C76"; CORAL = "EF6D5E"; MINT = "A6E4D5"
FOG = "E4ECE9"; PAPER = "F5F1E9"; MUTED = "586A6F"; WHITE = "FFFFFF"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_cell_text(cell, value, *, bold=False, color=INK, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(value)); set_font(r, "Aptos", size, color, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell_margins(cell)

def set_font(run, name="Aptos", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold; run.italic = italic

def add_field(paragraph, code):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = code
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])

def configure(doc, label):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(.78); sec.bottom_margin = Inches(.72)
    sec.left_margin = Inches(.82); sec.right_margin = Inches(.82)
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "Aptos"; normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.12
    for name,size,color,before,after in [
        ("Title",30,INK,0,8),("Subtitle",13,MUTED,0,18),
        ("Heading 1",20,INK,18,8),("Heading 2",13.5,TEAL,14,5),("Heading 3",11,INK,10,3)
    ]:
        st=styles[name]; st.font.name="Aptos Display" if name in ("Title","Heading 1") else "Aptos"
        st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color)
        st.font.bold=name not in ("Title","Subtitle")
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
        st.paragraph_format.keep_with_next=True
    header = sec.header
    p=header.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run("INFLECTION RADAR"); set_font(r,"Aptos",8,TEAL,True)
    r=p.add_run("   /   "+label.upper()); set_font(r,"Aptos",8,MUTED,False)
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr"); bottom=OxmlElement("w:bottom")
    for k,v in (("val","single"),("sz","5"),("space","5"),("color","C9D5D1")): bottom.set(qn(f"w:{k}"),v)
    pbdr.append(bottom); pPr.append(pbdr)
    footer=sec.footer; table=footer.add_table(1,2,Inches(6.86)); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width=Inches(5.8); table.columns[1].width=Inches(1.06)
    set_cell_text(table.cell(0,0),"Prepared for Will Hargreaves / Real Chemistry",color=MUTED,size=7.5)
    p=table.cell(0,1).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p.add_run("Page "); set_font(r,"Aptos",7.5,MUTED); add_field(p,"PAGE")
    table.cell(0,0)._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))
    core=doc.core_properties; core.author="Rayven-Nikkita Collins"
    core.subject="Inflection Radar"; core.keywords="Real Chemistry; Pilot 1; opportunity intelligence"

def p(doc,text,bold_lead=None,style=None):
    para=doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r=para.add_run(bold_lead); set_font(r,bold=True)
        r=para.add_run(text[len(bold_lead):]); set_font(r)
    else:
        r=para.add_run(text); set_font(r)
    return para

def bullet(doc,text):
    para=doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent=Inches(.28); para.paragraph_format.first_line_indent=Inches(-.18)
    para.paragraph_format.space_after=Pt(3)
    r=para.add_run(text); set_font(r)
    return para

def table(doc, headers, rows, widths=None, font=8.4):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=False
    if widths:
        for i,w in enumerate(widths): t.columns[i].width=Inches(w)
    for i,h in enumerate(headers):
        shade(t.cell(0,i),INK); set_cell_text(t.cell(0,i),h,bold=True,color=WHITE,size=font)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            shade(cells[i],WHITE if ri%2==0 else FOG); set_cell_text(cells[i],val,size=font)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def title_page(doc,title,subtitle,label):
    doc.add_paragraph()
    q=doc.add_paragraph(); q.paragraph_format.space_before=Pt(80); q.paragraph_format.space_after=Pt(18)
    r=q.add_run(label.upper()); set_font(r,"Aptos",9,TEAL,True)
    p=doc.add_paragraph(style="Title"); r=p.add_run(title); set_font(r,"Aptos Display",30,INK,False)
    p=doc.add_paragraph(style="Subtitle"); r=p.add_run(subtitle); set_font(r,"Aptos",13,MUTED)
    line=doc.add_paragraph(); line.paragraph_format.space_after=Pt(20)
    pPr=line._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr"); bottom=OxmlElement("w:bottom")
    for k,v in (("val","single"),("sz","18"),("space","1"),("color",CORAL)): bottom.set(qn(f"w:{k}"),v)
    pbdr.append(bottom); pPr.append(pbdr)
    table(doc,["Prepared for","Prepared by","Version"],[["Will Hargreaves / Real Chemistry","Rayven-Nikkita Collins","28 July 2026 · v2.0"]],[2.2,2.5,2.0],8.2)
    q=doc.add_paragraph(); q.paragraph_format.space_before=Pt(24); q.paragraph_format.left_indent=Inches(.25)
    q.paragraph_format.right_indent=Inches(.5)
    r=q.add_run("The public evidence establishes what Real Chemistry already demonstrates. It does not resolve what only Real Chemistry can know. Pilot 1 is designed around that boundary.")
    set_font(r,"Aptos Display",15,INK,True)
    doc.add_page_break()

def add_callout(doc,label,text,fill=FOG):
    t=doc.add_table(1,2); t.autofit=False
    t.columns[0].width=Inches(1.3); t.columns[1].width=Inches(5.3)
    shade(t.cell(0,0),fill); shade(t.cell(0,1),fill)
    set_cell_text(t.cell(0,0),label.upper(),bold=True,color=TEAL,size=8)
    set_cell_text(t.cell(0,1),text,size=10)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def build_brief():
    doc=Document(); configure(doc,"Executive decision brief")
    title_page(doc,"Inflection Radar","Executive decision brief for an eight-week emerging neuropsychiatry opportunity-intelligence pilot","Real Chemistry / internal discussion")
    doc.add_heading("The decision on the table",1)
    p(doc,"Approve the scope and minimal inputs for Pilot 1. No production integration, internal data, CRM connection, licensed Real Chemistry database, or ANATOMI build is required.")
    add_callout(doc,"Decision requested","Proceed with a 50–75-company public-source research universe and a 12–15-account human-reviewed priority portfolio.")
    doc.add_heading("What approval does not include",2)
    for x in ["A production AI product or autonomous target-selection system.","Use of client-confidential, patient-level, licensed-restricted, or internal Real Chemistry data.","ANATOMI, CRM, or production engineering integration.","A claim that Real Chemistry lacks AI, analytics, intelligence, monitoring, or growth capability."]: bullet(doc,x)
    doc.add_heading("The precise candidate white space",2)
    p(doc,"Real Chemistry publicly demonstrates nearly all technical and analytical ingredients. What is not publicly demonstrated is a unified, owned growth workflow that organizes early external company evidence around prospective-account decisions.")
    doc.add_heading("Will’s diligence questions: direct answers",1)
    table(doc,["Question","Answer"],[
        ["AI model","A governed model-task portfolio, not one model for every task. AI may assist bounded extraction, comparison, classification, contradiction detection, and draft synthesis. Humans retain consequential authority."],
        ["Assumptions","Every material assumption is registered and labeled evidence-supported, methodological, provisional, or internal-calibration-required."],
        ["Sources and permissions","Pilot 1 defaults to official, first-party, public-domain, and properly permitted sources. Rights and restrictions travel with every record and derived output."],
        ["Compliance","Company-level research; no PHI. Restricted sources are excluded by default. Human gates, audit records, and escalation pathways preserve healthcare, privacy, security, and AI-governance boundaries."],
        ["Reliability and limits","Facts, findings, inferences, and hypotheses remain distinct. Public evidence can support review priority; it cannot establish clients, conflicts, relationships, economics, capacity, or pursuit history."]
    ],[1.35,5.35],8.2)
    doc.add_heading("What Pilot 1 produces",1)
    table(doc,["Scale","Outputs","Boundary"],[
        ["8 weeks","Company universe; 12–15 priority dossiers; external opportunity score; rejected/deferred list; recommendation.","No production integration."],
        ["50–75 companies","Source-permission, assumption, evidence, and model-task registers.","No internal or client-confidential data."],
        ["Human-reviewed portfolio","Advance, monitor, defer, exclude, redirect, or insufficient-evidence dispositions.","No automated pursuit decision."]
    ],[1.1,3.35,2.25],8.2)
    add_callout(doc,"External question","Which companies appear externally attractive for Real Chemistry review?")
    add_callout(doc,"Internal question","Which externally attractive companies are plausible internal pursuits?",fill="F7E9E5")
    doc.add_heading("Three-stage progression",1)
    table(doc,["Stage","RN builds","Real Chemistry contributes","Decision"],[
        ["Discover / Pilot 1","Universe, ontology, evidence architecture, external score, dossiers, registers.","Scope, geography, stage, reader, public priorities, success standard.","Which companies warrant review?"],
        ["Qualify / conditional","Internal viability layer and calibrated review workflow.","Clients, pursuits, conflicts, ownership, thresholds, economics, judgment.","Which are plausible pursuits?"],
        ["Institutionalize / conditional","Monitoring, queues, outcome capture, recalibration design.","Approved systems, CRM, data feeds, ownership decisions.","Should this become persistent?"]
    ],[1.15,1.85,1.9,1.8],7.6)
    doc.add_heading("Eight-week cadence",2)
    table(doc,[str(i) for i in range(1,9)],[["Scope","Taxonomy","Universe","Verify","Score","Dossiers","Portfolio","Handoff"]],[.84]*8,7.2)
    doc.add_heading("Safeguards and evaluation",1)
    for x in ["Five human gates: source approval, record validation, analytical approval, account review, and distribution approval.","Four evidence tiers: verified fact, corroborated finding, analytical inference, and hypothesis or early signal.","Score sensitivity, missing-data review, contradiction review, staleness checks, baseline comparison, and false-positive analysis.","Success means a defensible and useful internal review portfolio—not a promise that every ranked company becomes an account."]: bullet(doc,x)
    doc.add_heading("What RN needs from Will now",1)
    table(doc,["Input","Why it matters"],[
        ["Scope","Defines the therapeutic and company universe."],["Intended reader","Determines depth, language, and format."],
        ["Public service priorities","Constrains capability mapping without relying on internal data."],["Success standard","Defines what a useful portfolio must enable."]
    ],[2.0,4.7],8.5)
    add_callout(doc,"Decision","Approve Discover/Pilot 1 first. Reserve internal calibration, ANATOMI, CRM, licensed-source enrichment, and production integration for a later decision.")
    doc.save(ASSETS/"executive-brief.docx")

def build_report():
    doc=Document(); configure(doc,"Internal research report")
    title_page(doc,"Inflection Radar","Emerging Neuropsychiatry Opportunity Intelligence · Internal research report","Real Chemistry / internal discussion")
    doc.add_heading("Executive summary",1)
    p(doc,"Pilot 1 is an eight-week, public-source research engagement covering a 50–75-company universe and producing a human-reviewed portfolio of 12–15 priority accounts. It can be completed without ANATOMI, CRM, internal Real Chemistry data, licensed Real Chemistry databases, or production engineering.")
    p(doc,"The defensible white-space conclusion is narrow: Real Chemistry publicly demonstrates nearly all technical and analytical ingredients required for an emerging-company opportunity system. The public record does not establish a unified, owned growth workflow organized around the upstream decision of which emerging companies merit attention now, why, and which practice should inspect the opportunity.")
    add_callout(doc,"Protected distinction","Pilot 1 asks which companies appear externally attractive for review. A later internal-calibration phase asks which are plausible internal pursuits.")
    doc.add_heading("Questions presented",1)
    for q in ["Is the proposed pilot technically feasible and non-duplicative?","What exactly is Pilot 1, and what does RN build independently?","What must Real Chemistry provide now versus later?","How are sources, model assumptions, permissions, compliance, reliability, and limitations governed?","How would the method progress from research to internal calibration to an owned capability?","What would Real Chemistry receive, and how should usefulness be evaluated?"]: bullet(doc,q)
    doc.add_heading("Conclusions",1)
    table(doc,["Issue","Conclusion"],[
        ["Feasibility","Yes. Pilot 1 can operate in an open-research deployment zone using approved public and open sources."],
        ["Duplication","The method must complement—not reproduce—ANATOMI, Analytics & Insights, HealthGEO, ReputAI, RC Resolve, or other existing capabilities."],
        ["AI","Use a governed model-task portfolio. Provider-flexible does not mean model-indifferent."],
        ["Authority","Humans retain company inclusion, interpretation, scoring override, prioritization, and distribution authority."],
        ["Limit","Public evidence cannot resolve internal clients, pursuits, conflicts, relationships, economics, ownership, or capacity."]
    ],[1.25,5.45],8.4)
    doc.add_heading("1. Pilot 1: complete initial engagement",1)
    table(doc,["Dimension","Specification"],[
        ["Duration","Eight weeks"],["Universe","50–75 companies"],["Portfolio","12–15 human-reviewed priority accounts"],
        ["Records","Source-linked dossiers, evidence ledger, assumption register, model-task register, source-permission register"],
        ["Decision outputs","External opportunity score; advance/monitor/defer/exclude/redirect/insufficient-evidence dispositions; rejected/deferred list; recommendation"],
        ["Excluded","Production integration, internal data, CRM, PHI, restricted sources, automated pursuit decisions"]
    ],[1.45,5.25],8.4)
    doc.add_heading("Three-stage progression",2)
    table(doc,["Stage","Question","Additional inputs","System state"],[
        ["Discover","Which companies warrant review?","Minimal public-scope inputs","Independent open-research pilot"],
        ["Qualify","Which are plausible pursuits?","Clients, pursuits, conflicts, ownership, economics, thresholds","Controlled internal calibration"],
        ["Institutionalize","Should this become persistent?","Approved systems, CRM, monitoring, ownership, outcomes","Real Chemistry-owned capability"]
    ],[1.1,1.8,2.2,1.6],7.7)
    doc.add_heading("2. Evidence and reasoning architecture",1)
    p(doc,"The unit of work is not a generated narrative. It is a traceable chain from source to factual record, normalized signal, explicit hypothesis, possible capability relevance, score rationale, human review, and disposition.")
    table(doc,["Stage","Required record","Authority"],[
        ["Source","URL/title/publisher/access date/rights class/access method/restrictions","Source approval"],
        ["Fact","Entity-resolved, dated, source-grounded record","Record validation"],
        ["Signal","Defined ontology family and observed change","Analytical approval"],
        ["Hypothesis","Why the signal may matter; alternatives and contradictions","Account review"],
        ["Decision","Advance, monitor, defer, exclude, redirect, or insufficient evidence","Distribution approval"]
    ],[1.0,3.9,1.8],8.0)
    doc.add_heading("Evidence tiers",2)
    table(doc,["Tier","Meaning"],[
        ["Verified fact","Directly established by an authoritative or primary record."],
        ["Corroborated finding","Supported by multiple reliable records or validated against an authoritative record."],
        ["Analytical inference","A reasoned interpretation that remains distinct from its factual foundation."],
        ["Hypothesis or early signal","A proposition requiring further evidence or internal review."]
    ],[1.7,5.0],8.4)
    doc.add_heading("3. Model-task portfolio",1)
    p(doc,"The pilot does not depend on one undifferentiated model. Model selection is task-specific and conditioned on sensitivity, source restrictions, provider terms, retention, training use, accuracy, observability, security, cost, performance, and suitability.")
    table(doc,["Lane","Permitted work","Authority boundary"],[
        ["AI-assisted","Extraction, entity resolution, comparison, classification, contradiction detection, and draft synthesis.","May propose; cannot own consequential decisions."],
        ["Deterministic","Required fields, record IDs, registry checks, score calculation, recency rules, versioning, and citations.","Rules must be inspectable and versioned."],
        ["Human-only","Inclusion, interpretation, scoring overrides, escalation, prioritization, and distribution.","Named reviewer authority is mandatory."]
    ],[1.2,3.25,2.25],8.1)
    doc.add_heading("4. Assumption register",1)
    table(doc,["Assumption family","Required entries","Status labels"],[
        ["Scope","Therapeutic scope, geography, company stage, modalities","Evidence-supported / provisional"],
        ["Signals","Definition of inflection point; observable indications of possible service need","Methodological / provisional"],
        ["Mapping","How signal families map to Real Chemistry capabilities","Evidence-supported / internal calibration required"],
        ["Evidence","Thresholds, recency windows, missing-data treatment","Methodological"],
        ["Scoring","Weights, false-positive risk, sensitivity boundaries","Provisional / internal calibration required"],
        ["Limits","What public evidence cannot establish","Evidence-supported"]
    ],[1.3,3.5,1.9],8.0)
    doc.add_heading("5. Scoring and portfolio decision",1)
    p(doc,"No single score is the final answer. The method separates an externally observable score from internal pursuit viability and the final human portfolio disposition.")
    table(doc,["Layer","Inputs","Output"],[
        ["External Opportunity Score","Strategic relevance, timing, evidence strength, service-need clarity, commercialization complexity, narrative/reputation complexity, differentiation, risk-adjusted review priority","Externally grounded review priority"],
        ["Internal Pursuit Viability","Client status, active pursuit, conflict, relationship, history, economics, capacity, ownership, geography, competitive position, strategic priority","Real Chemistry-calibrated viability"],
        ["Human Portfolio Decision","Evidence, scores, contradictions, reviewer judgment, unresolved questions","Advance / monitor / defer / exclude / redirect / insufficient evidence"]
    ],[1.45,3.6,1.65],7.8)
    doc.add_heading("6. Source-permission architecture",1)
    p(doc,"Permission is evaluated across the full lifecycle. A source label alone is insufficient, and derived intelligence inherits upstream restrictions.")
    table(doc,["Architecture","Elements"],[
        ["8 lifecycle layers","Decision/use case; acquisition; rights; engineering/security; retrieval/provenance; analysis/orchestration; validation/compliance; output/action/monitoring"],
        ["7 rights classes","Government/public domain; public copyrighted; licensed commercial; Real Chemistry proprietary; client confidential; sensitive/regulated; derived intelligence"],
        ["5 data states","Raw; normalized; indexed; derived; distributed"],
        ["8 control planes","Purpose/authority; rights/data protection; security; model/agent governance; healthcare compliance; evidence integrity; accountability; monitoring/incident response"],
        ["4 deployment zones","Open research; controlled enterprise; client-restricted; regulated health/safety"],
        ["3 incident pathways","Security/privacy; AI/evidence integrity; healthcare safety"]
    ],[1.45,5.25],8.0)
    doc.add_heading("Pilot 1 source perimeter",2)
    table(doc,["Core","Conditional on rights/approval","Excluded initially"],[
        ["SEC; FDA; ClinicalTrials.gov; USPTO; official registries; official company sites; first-party releases; official agendas; PubMed/Crossref metadata; properly licensed open-access research","Licensed news; commercial biotech databases; PitchBook; social listening; publisher full text; conference slides/recordings; LinkedIn; job boards","Scraped LinkedIn/job boards; attendee lists; PHI; inferred health data; purchased contacts; unapproved confidential material; access-controlled sources without permission; full commercial/publisher corpora without rights"]
    ],[2.2,2.2,2.3],7.6)
    doc.add_heading("7. Real Chemistry capability fit",1)
    p(doc,"The proposal should not imply absence where public evidence is silent. Every capability finding is classified as Confirmed Have, Confirmed Boundary, Not Publicly Evidenced, or Internally Unknown.")
    table(doc,["Capability","Class","Publicly established","Internal question"],[
        ["ANATOMI","Confirmed have","AI ecosystem combining data, models, orchestration, analytics, and agentic tools.","Is it used for internal growth or prospect discovery?"],
        ["Analytics & Insights","Confirmed have","Named intelligence, listening, audience, prediction, and measurement capabilities.","Which products already support prospective-account decisions?"],
        ["HealthGEO / ReputAI","Confirmed have","AI-search perception, source intelligence, reputation measurement, and predictive message testing.","Can outputs inform early account hypotheses?"],
        ["Unified prospect workflow","Not publicly evidenced","No public evidence located for the specific source-to-growth-decision workflow proposed here.","Does an equivalent internal process already exist?"],
        ["CRM, conflicts, ownership, outcomes","Internally unknown","Public evidence cannot resolve the systems, rules, or history.","What can be safely used during calibration?"]
    ],[1.2,1.0,2.6,1.9],7.4)
    doc.add_heading("8. Eight-week execution plan",1)
    table(doc,["Week","Activities","Output","Gate"],[
        ["1","Scope, decision, perimeter, governance, schema","Scope memo + source perimeter","Scope approval"],
        ["2","Taxonomy, assumptions, evidence rules, entity model","Ontology + registers","Method approval"],
        ["3","Discovery, universe, inclusion/exclusion","50–75-company universe","Universe review"],
        ["4","Primary-source verification, normalized records, entity resolution","Evidence records","Record validation"],
        ["5","Scoring, sensitivity, false positives, missing data","Score + sensitivity memo","Analytical approval"],
        ["6","Dossier production","Draft priority dossiers","Account review"],
        ["7","Portfolio and rejected/deferred list","12–15-account portfolio","Portfolio review"],
        ["8","Evaluation, handoff, recommendation, future-stage decision","Final package","Distribution approval"]
    ],[.55,2.75,2.2,1.2],7.4)
    doc.add_heading("9. Dossier specification",1)
    for x in ["Account snapshot","Evidence summary","Clinical and regulatory position","Corporate movement","Commercialization-readiness signals","Narrative and reputation landscape","Real Chemistry relevance","Source and permission record","Human-review notes","Suggested internal next step"]: bullet(doc,x)
    doc.add_heading("10. Evaluation framework",1)
    table(doc,["Dimension","Checks"],[
        ["Research quality","Citation accuracy, coverage, staleness, contradiction handling, entity resolution"],
        ["Decision usefulness","False positives, rationale quality, reviewer usefulness, disposition clarity"],
        ["Commercial relevance","Review-worthy accounts, service-fit clarity, timing insight"],
        ["Governance","Permission completeness, restricted-source exclusion, audit completeness, escalation"],
        ["Learning","Whether signal families later correlate with interest, deferral, pursuit, rejection, or win/loss"]
    ],[1.35,5.35],8.2)
    doc.add_heading("11. Internal questions for Will",1)
    questions={
        "Current process":["Who identifies entirely new prospective companies today?","Are target lists centralized or practice-specific?","Where are early company signals recorded?"],
        "Existing capability":["Does Analytics & Insights support internal growth?","Is ANATOMI used for prospect discovery or qualification?","Is there already a target-account score, dashboard, or watchlist?"],
        "Internal calibration":["How are clients, pursuits, and conflicts screened?","Who owns a prospective account before a formal opportunity exists?","What economics, capacity, and thresholds matter?"],
        "Technology":["Which AI environment is approved?","Can a structured company dataset and ontology be accepted?","Which licensed sources and growth CRM are available?"],
        "Measurement":["What would make the portfolio useful or duplicative?","Are historical examples and structured win/loss reasons available?","Could outcomes be connected to originating signals?"]
    }
    for head,items in questions.items():
        doc.add_heading(head,2)
        for x in items: bullet(doc,x)
    doc.add_heading("Appendix A. Minimum data structure",1)
    table(doc,["Entity","Minimum content"],[
        ["Company","Canonical identity and aliases"],["Source","URL, title, publisher, access date, rights class"],["Fact","Source-grounded factual record"],["Signal","Ontology family and normalized event"],["Evidence tier","Verified, corroborated, inference, or hypothesis"],["Inflection hypothesis","Why the signal may matter now"],["Service relevance","Possible capability fit"],["Score","Dimensions, rationale, and weight version"],["Review","Reviewer, disposition, override, unresolved question"],["Governance","Restrictions, audience, retention, approvals"]
    ],[1.4,5.3],8.2)
    doc.add_heading("Appendix B. Verification and source note",1)
    p(doc,"The capability assessment relies on public Real Chemistry materials and official public sources. Dynamic webpages, product names, access conditions, and technical descriptions must be rechecked immediately before circulation. Public-source classifications are not representations about non-public Real Chemistry systems.")
    table(doc,["Publisher","Use","URL"],[
        ["Real Chemistry","ANATOMI, Analytics & Insights, HealthGEO, ReputAI, and responsible-healthcare-AI materials","https://www.realchemistry.com/"],
        ["U.S. SEC","Company filings","https://www.sec.gov/edgar/search/"],["U.S. FDA","Regulatory records","https://www.fda.gov/"],
        ["ClinicalTrials.gov","Trial records","https://clinicaltrials.gov/"],["USPTO","Patent and trademark records","https://www.uspto.gov/"],
        ["PubMed","Scholarly metadata","https://pubmed.ncbi.nlm.nih.gov/"],["Crossref","Scholarly metadata","https://www.crossref.org/"]
    ],[1.25,3.0,2.45],7.5)
    doc.save(ASSETS/"full-report.docx")

if __name__=="__main__":
    build_brief(); build_report()
    print(ASSETS/"executive-brief.docx")
    print(ASSETS/"full-report.docx")
